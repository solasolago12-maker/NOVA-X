"""File export utilities for NOVA-X.

Supports exporting content to multiple formats: plain text, Markdown,
Microsoft Word (.docx), and PDF.  If optional dependencies are missing,
exports gracefully fall back to plain text.
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Optional


class FileExporter:
    """Multi-format file exporter.

    Args:
        output_dir: Directory where files are written. Defaults to
            ``~/Documents/NOVA-X``.
    """

    def __init__(self, output_dir: Optional[str] = None) -> None:
        self.output_dir = Path(output_dir) if output_dir else Path.home() / "Documents" / "NOVA-X"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------------

    def export_txt(self, content: str, filename: Optional[str] = None) -> str:
        """Export content as a plain text file.

        Args:
            content: Text to export.
            filename: Optional filename (``.txt`` appended if missing).

        Returns:
            Absolute path to the written file.
        """
        name = filename or f"nova_x_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        if not name.endswith(".txt"):
            name += ".txt"
        filepath = self.output_dir / name
        with open(filepath, "w", encoding="utf-8") as fh:
            fh.write(content)
        return str(filepath)

    # ------------------------------------------------------------------
    # Markdown
    # ------------------------------------------------------------------

    def export_md(self, content: str, title: str = "NOVA-X Export", filename: Optional[str] = None) -> str:
        """Export content as a Markdown file with front-matter.

        Args:
            content: Markdown body text.
            title: Document title for front-matter.
            filename: Optional filename (``.md`` appended if missing).

        Returns:
            Absolute path to the written file.
        """
        name = filename or f"nova_x_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        if not name.endswith(".md"):
            name += ".md"
        filepath = self.output_dir / name
        header = f"---\ntitle: {title}\ndate: {datetime.now().isoformat()}\n---\n\n"
        with open(filepath, "w", encoding="utf-8") as fh:
            fh.write(header + content)
        return str(filepath)

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def export_docx(self, content: str, title: str = "NOVA-X Export", filename: Optional[str] = None) -> str:
        """Export content as a Word document.

        Falls back to ``.txt`` if ``python-docx`` is not installed.

        Args:
            content: Text content to export.
            title: Document title.
            filename: Optional filename (``.docx`` appended if missing).

        Returns:
            Absolute path to the written file.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("[WARNING] python-docx not installed. Falling back to .txt export.")
            return self.export_txt(content, filename)

        name = filename or f"nova_x_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        if not name.endswith(".docx"):
            name += ".docx"
        filepath = self.output_dir / name

        doc = Document()

        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        date_run.italic = True
        date_run.font.size = Pt(10)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer

        # Content
        for line in content.split("\n"):
            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("1. ") or line.startswith("2. "):
                doc.add_paragraph(line[3:], style="List Number")
            elif line.strip() == "":
                doc.add_paragraph()
            else:
                doc.add_paragraph(line)

        doc.save(filepath)
        return str(filepath)

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def export_pdf(self, content: str, title: str = "NOVA-X Export", filename: Optional[str] = None) -> str:
        """Export content as a PDF document.

        Falls back to ``.txt`` if ``reportlab`` is not installed.

        Args:
            content: Text content to export.
            title: Document title.
            filename: Optional filename (``.pdf`` appended if missing).

        Returns:
            Absolute path to the written file.
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_CENTER
        except ImportError:
            print("[WARNING] reportlab not installed. Falling back to .txt export.")
            return self.export_txt(content, filename)

        name = filename or f"nova_x_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        if not name.endswith(".pdf"):
            name += ".pdf"
        filepath = self.output_dir / name

        doc = SimpleDocTemplate(str(filepath), pagesize=letter,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)

        styles = getSampleStyleSheet()
        story: list = []

        # Title
        title_style = ParagraphStyle("CustomTitle", parent=styles["Heading1"],
                                     alignment=TA_CENTER, spaceAfter=12)
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(
            f"<i>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</i>",
            ParagraphStyle("Date", parent=styles["Normal"], alignment=TA_CENTER, spaceAfter=20)
        ))
        story.append(Spacer(1, 0.2 * inch))

        # Body
        normal = styles["Normal"]
        for line in content.split("\n"):
            if line.strip() == "":
                story.append(Spacer(1, 0.1 * inch))
            elif line.startswith("# "):
                story.append(Paragraph(line[2:], styles["Heading1"]))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], styles["Heading2"]))
            elif line.startswith("### "):
                story.append(Paragraph(line[4:], styles["Heading3"]))
            elif line.startswith("- "):
                story.append(Paragraph(f"&bull; {line[2:]}", normal))
            else:
                story.append(Paragraph(line, normal))

        doc.build(story)
        return str(filepath)
